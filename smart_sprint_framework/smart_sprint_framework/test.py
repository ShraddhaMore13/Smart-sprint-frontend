# test_docx.py
import docx
import os

def test_docx():
    # Create a simple test document
    doc = docx.Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph.')
    
    test_path = 'test_document.docx'
    doc.save(test_path)
    
    # Try to read it back
    try:
        doc = docx.Document(test_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        content = '\n'.join(text)
        print("Document content:")
        print(content)
        
        # Clean up
        os.remove(test_path)
        print("Test passed! python-docx is working correctly.")
        return True
    except Exception as e:
        print(f"Test failed: {e}")
        return False

if __name__ == "__main__":
    test_docx()