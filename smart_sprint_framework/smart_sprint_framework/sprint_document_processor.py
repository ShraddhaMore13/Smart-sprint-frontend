import re
import docx
import csv
import os
from nlp_pipeline import NLPPipeline
import json

class SprintDocumentProcessor:
    def __init__(self):
        self.nlp = NLPPipeline()
        
        # Define task templates for common user story patterns
        self.task_templates = {
            "password reset": [
                "Design password reset flow",
                "Implement password reset request form",
                "Create secure token generation",
                "Implement email notification system",
                "Develop password update page",
                "Test password reset functionality"
            ],
            "authentication": [
                "Design login form",
                "Implement authentication logic",
                "Create user session management",
                "Develop logout functionality",
                "Test authentication flow"
            ],
            "user profile": [
                "Design profile page layout",
                "Implement profile data retrieval",
                "Create profile update functionality",
                "Implement profile image upload",
                "Test profile management"
            ],
            "registration": [
                "Design registration form",
                "Implement user registration logic",
                "Create email verification system",
                "Develop registration confirmation page",
                "Test registration process"
            ],
            "email sending": [
                "Design email template",
                "Implement email service integration",
                "Create email queue system",
                "Add email tracking",
                "Test email delivery"
            ],
            "password reset page": [
                "Design password reset page layout",
                "Implement secure token validation",
                "Create password update form",
                "Add password strength validation",
                "Test page functionality"
            ],
            "password update": [
                "Design password update logic",
                "Implement password hashing",
                "Create password validation",
                "Update user authentication",
                "Test password update functionality"
            ],
            "confirmation page": [
                "Design confirmation page layout",
                "Implement success message display",
                "Add redirect to login",
                "Create user notification",
                "Test confirmation flow"
            ],
            "test password reset": [
                "Create test plan for password reset",
                "Test password reset request",
                "Test email notification",
                "Test password update",
                "Test security aspects",
                "Document test results"
            ],
            "default": [
                "Analyze requirements",
                "Design solution",
                "Implement core functionality",
                "Create unit tests",
                "Test integration",
                "Document implementation"
            ]
        }
    
    def process_document(self, doc_path):
        """Process a Word document and extract multiple tasks"""
        try:
            doc = docx.Document(doc_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            document_text = '\n'.join(full_text)
            
            # Split the document into tasks using regex
            # Pattern: "Task X:" or "Task X" where X is a number
            task_pattern = r'Task\s+\d+[:\s]*(.*?)(?=Task\s+\d+[:\s]*|$)'
            matches = re.findall(task_pattern, document_text, re.DOTALL)
            
            tasks = []
            for i, match in enumerate(matches):
                # Each match is the content of one task
                task_text = match.strip()
                
                # Extract title: the first line of the task_text might be the title
                lines = task_text.split('\n')
                title = lines[0].strip()
                # The rest is the description
                description = '\n'.join(lines[1:]).strip()
                
                # If the title is empty, generate one
                if not title:
                    title = f"Task {i+1}"
                
                # Extract priority from the task_text using the existing NLP method
                entities = self.nlp.extract_entities(task_text)
                priority = entities['priorities'][0] if entities['priorities'] else 'medium'
                
                # Estimate hours based on complexity of the task description
                complexity = self.nlp.analyze_complexity(description)
                estimated_hours = complexity * 8  # Simple heuristic
                
                tasks.append({
                    'title': title,
                    'description': description,
                    'priority': priority,
                    'estimated_hours': estimated_hours
                })
            
            return tasks
        
        except Exception as e:
            print(f"Error processing document: {e}")
            return []
    
    def process_text_document(self, txt_path):
        """Process a text document and extract tasks"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                document_text = file.read()
            
            # Split the document into tasks using regex
            # Pattern: "Task X:" or "Task X" where X is a number
            task_pattern = r'Task\s+\d+[:\s]*(.*?)(?=Task\s+\d+[:\s]*|$)'
            matches = re.findall(task_pattern, document_text, re.DOTALL)
            
            tasks = []
            for i, match in enumerate(matches):
                # Each match is the content of one task
                task_text = match.strip()
                
                # Extract title: the first line of the task_text might be the title
                lines = task_text.split('\n')
                title = lines[0].strip()
                # The rest is the description
                description = '\n'.join(lines[1:]).strip()
                
                # If the title is empty, generate one
                if not title:
                    title = f"Task {i+1}"
                
                # Extract priority from the task_text using the existing NLP method
                entities = self.nlp.extract_entities(task_text)
                priority = entities['priorities'][0] if entities['priorities'] else 'medium'
                
                # Estimate hours based on complexity of the task description
                complexity = self.nlp.analyze_complexity(description)
                estimated_hours = complexity * 8  # Simple heuristic
                
                tasks.append({
                    'title': title,
                    'description': description,
                    'priority': priority,
                    'estimated_hours': estimated_hours
                })
            
            return tasks
        
        except Exception as e:
            print(f"Error processing text document: {e}")
            return []
    
    def process_sprint_document_with_stories(self, doc_path):
        """Process a sprint document with sprint goal and user stories, and generate tasks"""
        try:
            # Read the document
            if doc_path.lower().endswith('.docx'):
                doc = docx.Document(doc_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                document_text = '\n'.join(full_text)
            elif doc_path.lower().endswith('.txt'):
                with open(doc_path, 'r', encoding='utf-8') as file:
                    document_text = file.read()
            else:
                print("Error: Only .docx and .txt files are supported")
                return None
            
            # Extract sprint goal
            sprint_goal_match = re.search(r'Sprint Goal:\s*"([^"]*)"', document_text)
            sprint_goal = sprint_goal_match.group(1) if sprint_goal_match else "No sprint goal specified"
            
            # Extract user stories
            user_stories = []
            user_story_pattern = r'"As a[^"]*"\s*\([^)]*\)'
            user_story_matches = re.findall(user_story_pattern, document_text)
            
            for match in user_story_matches:
                # Extract the story and story points
                story_match = re.match(r'"(As a[^"]*)"\s*\([^)]*Story Points:\s*(\d+)[^)]*\)', match)
                if story_match:
                    story = story_match.group(1)
                    story_points = int(story_match.group(2))
                    user_stories.append({
                        'story': story,
                        'story_points': story_points
                    })
            
            # Determine the feature type from all user stories
            feature_type = None
            for story in user_stories:
                current_feature_type = self._extract_feature_type(story['story'])
                # If we haven't determined a feature type yet, use the first one we find
                if feature_type is None:
                    feature_type = current_feature_type
            
            # If no specific feature type was found, use default
            if feature_type is None:
                feature_type = 'default'
            
            # Get task templates for this feature type
            templates = self.task_templates.get(feature_type, self.task_templates['default'])
            
            # Generate tasks from templates (only once, not per story)
            tasks = []
            for template in templates:
                # Calculate estimated hours based on average story points and task complexity
                if user_stories:
                    avg_story_points = sum(story['story_points'] for story in user_stories) / len(user_stories)
                else:
                    avg_story_points = 3  # Default if no stories found
                
                task_complexity = self._analyze_task_complexity(template)
                estimated_hours = max(2, avg_story_points * task_complexity)
                
                # Round up to nearest integer
                estimated_hours = int(estimated_hours) + (1 if estimated_hours % 1 > 0 else 0)
                
                # Determine priority based on task complexity
                priority = 'high' if task_complexity > 1.5 else 'medium' if task_complexity > 1.0 else 'low'
                
                tasks.append({
                    'title': template,
                    'description': f"Task for sprint: {sprint_goal}",
                    'priority': priority,
                    'estimated_hours': estimated_hours,
                    'story_points': avg_story_points
                })
            
            return {
                'sprint_goal': sprint_goal,
                'user_stories': user_stories,
                'tasks': tasks
            }
            
        except Exception as e:
            print(f"Error processing sprint document: {e}")
            return None
    
    def _extract_feature_type(self, story):
        """Extract the feature type from a user story"""
        story_lower = story.lower()
        
        # Check for specific feature keywords
        if 'password reset' in story_lower:
            return 'password reset'
        elif 'authentication' in story_lower or 'login' in story_lower or 'log in' in story_lower:
            return 'authentication'
        elif 'profile' in story_lower:
            return 'user profile'
        elif 'registration' in story_lower or 'register' in story_lower or 'sign up' in story_lower:
            return 'registration'
        elif 'email sending' in story_lower:
            return 'email sending'
        elif 'password reset page' in story_lower:
            return 'password reset page'
        elif 'password update' in story_lower:
            return 'password update'
        elif 'confirmation page' in story_lower:
            return 'confirmation page'
        elif 'test password reset' in story_lower:
            return 'test password reset'
        else:
            return 'default'
    
    def _analyze_task_complexity(self, task_description):
        """Analyze the complexity of a task description"""
        # Simple heuristic based on keywords
        task_lower = task_description.lower()
        
        # High complexity tasks
        high_complexity_keywords = ['implement', 'develop', 'create', 'design', 'integrate']
        # Medium complexity tasks
        medium_complexity_keywords = ['test', 'update', 'modify', 'improve']
        # Low complexity tasks
        low_complexity_keywords = ['document', 'review', 'check']
        
        complexity = 1.0  # Default complexity
        
        for keyword in high_complexity_keywords:
            if keyword in task_lower:
                complexity += 0.5
        
        for keyword in medium_complexity_keywords:
            if keyword in task_lower:
                complexity += 0.3
        
        for keyword in low_complexity_keywords:
            if keyword in task_lower:
                complexity += 0.1
        
        return complexity
    
    def process_documents_to_csv(self, input_dir, output_csv):
        """Process all Word documents in a directory and save to CSV"""
        all_tasks = []
        
        # Process all Word documents in the directory
        doc_files = [f for f in os.listdir(input_dir) if f.endswith('.docx')]
        
        for doc_file in doc_files:
            doc_path = os.path.join(input_dir, doc_file)
            tasks = self.process_document(doc_path)
            all_tasks.extend(tasks)
        
        # Write to CSV
        with open(output_csv, 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = ['title', 'description', 'priority', 'estimated_hours']
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            
            writer.writeheader()
            for task in all_tasks:
                writer.writerow(task)
        
        return all_tasks