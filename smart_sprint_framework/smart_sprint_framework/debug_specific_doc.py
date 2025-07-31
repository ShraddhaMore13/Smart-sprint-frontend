# debug_specific_doc.py
import docx
import os
import re
import zipfile
import xml.etree.ElementTree as ET

def debug_specific_document():
    doc_path = r"C:\Users\shrad\source\repos\smart_sprint_framework\smart_sprint_framework\sprint_documents\Sprint 3 Features.docx"
    
    print(f"Debugging document: {doc_path}")
    
    # Check if file exists
    if not os.path.exists(doc_path):
        print("ERROR: File not found")
        return
    
    try:
        # Method 1: Try to open with python-docx
        print("\n=== Method 1: Using python-docx ===")
        try:
            doc = docx.Document(doc_path)
            print(f"Document opened successfully")
            print(f"Number of paragraphs: {len(doc.paragraphs)}")
            
            # Print all paragraphs
            print("\n--- Paragraph Content ---")
            for i, para in enumerate(doc.paragraphs):
                text = para.text
                print(f"Paragraph {i}: '{text}'")
            
            # Extract all text
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            document_text = '\n'.join(full_text)
            
            print(f"\n--- Full Document Text ---")
            print(document_text)
            
            # Look for task patterns
            task_pattern = r'Task\s+\d+[:\s]*(.*?)(?=Task\s+\d+[:\s]*|$)'
            matches = re.findall(task_pattern, document_text, re.DOTALL)
            
            print(f"\n--- Task Matches ---")
            print(f"Found {len(matches)} task matches")
            for i, match in enumerate(matches):
                print(f"Task {i+1}: '{match.strip()}'")
                
        except Exception as e:
            print(f"ERROR with python-docx: {e}")
            import traceback
            traceback.print_exc()
        
        # Method 2: Try to extract text using zipfile
        print("\n=== Method 2: Using zipfile to extract document.xml ===")
        try:
            with zipfile.ZipFile(doc_path) as zf:
                # List all files in the zip
                print("Files in the document:")
                for name in zf.namelist():
                    print(f"  {name}")
                
                # Try to read document.xml
                document_xml = zf.read('word/document.xml')
                print(f"\nSuccessfully read document.xml (size: {len(document_xml)} bytes)")
                
                # Parse XML
                root = ET.fromstring(document_xml)
                
                # Extract text from XML
                namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                text_nodes = root.findall('.//w:t', namespaces)
                
                text_parts = []
                for node in text_nodes:
                    if node.text:
                        text_parts.append(node.text)
                
                document_text = ''.join(text_parts)
                
                print(f"\n--- Extracted Text ---")
                print(document_text)
                
                # Look for task patterns
                task_pattern = r'Task\s+\d+[:\s]*(.*?)(?=Task\s+\d+[:\s]*|$)'
                matches = re.findall(task_pattern, document_text, re.DOTALL)
                
                print(f"\n--- Task Matches ---")
                print(f"Found {len(matches)} task matches")
                for i, match in enumerate(matches):
                    print(f"Task {i+1}: '{match.strip()}'")
                
                # Save extracted text to a file
                txt_path = os.path.splitext(doc_path)[0] + '_extracted.txt'
                with open(txt_path, 'w', encoding='utf-8') as f:
                    f.write(document_text)
                print(f"\nExtracted text saved to: {txt_path}")
                
        except Exception as e:
            print(f"ERROR with zipfile method: {e}")
            import traceback
            traceback.print_exc()
        
    except Exception as e:
        print(f"GENERAL ERROR: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    debug_specific_document()